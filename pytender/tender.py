from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path


# Start
firm = 'Shree Siddhababa Nirman Sewa'
firm_addr = 'Waling-9, Syangja'

# Ask if JV
jv = input('\nIf Joint Venture, write the full JV name, OTHERWISE skip this step: ')

while True:
    short_name = input('\nName of Contract in short: ')
    if short_name: break
path = short_name + '//SOURCE'
Path(path).mkdir(parents=True, exist_ok=True)

# Large or Small?
while True:
    ls = input('\nEnter 1 for Letter of Bid, and 2 for Letter of Technical and Price Bids: ')
    if ls == '1' or ls == '2': break


date = input('\nDate: ')
contract_name = input('\nName of the Contract: ')
IoB = input('\nInvitation for Bid No.: ')
Contract_ID = input('\nContract Identification No: ')

# Inside Address
lines = []
print('\n\nTo:')
while True:
    line = input()
    if line:
        lines.append(line)
    else:
        break
to = '\n'.join(lines)

days = input('\nEnter bid validity period: ')
if ls == '2':
    line_of_credit = input('\nEnter Line of Credit amount: ')
person = input('\nAuthorized Person: ')

class letter:
    def __init__(self):
        path_to_module = os.path.dirname(__file__)
        format_path = os.path.join(path_to_module, 'format.docx')
        self.doc = Document(format_path)
        
        
    def write_title(self, size=16):
        title = self.doc.add_paragraph(self.heading)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.runs[0].underline = title.runs[0].bold = True
        title.runs[0].font.size = Pt(size)


    def write_details(self):
        details = self.doc.add_paragraph('Date: ')
        details.alignment = WD_ALIGN_PARAGRAPH.RIGHT    
        details.add_run(date).bold = True
        details.add_run('\nName of the contract: ')    
        details.add_run(contract_name).bold = True
        details.add_run('\nInvitation for Bid No.: ')    
        details.add_run(IoB).bold = True
        details.add_run('\nContract Identification No.: ')
        details.add_run(Contract_ID).bold = True


    def write_inside_addr(self):
        inside_addr = self.doc.add_paragraph('To:\n')
        inside_addr.alignment = WD_ALIGN_PARAGRAPH.LEFT
        inside_addr.add_run(to).bold = True


    def write_signature(self):
        signature = self.doc.add_paragraph('\n\nName: ')
        signature.add_run(person).bold = True
        signature.add_run('\nIn the capacity of ')
        if jv:
            signature.add_run('Authorized JV Partner').bold = True
        else:
            signature.add_run('Proprietor').bold = True
        signature.add_run('\n\n\n\n\n\nSigned: ....................')
        signature.add_run('\nDuly authorized to sign the Bid for and on behalf of ')
        if jv:
            signature.add_run(jv).bold = True
        else:
            signature.add_run(firm).bold = True
        signature.add_run('\nDate: ')
        signature.add_run(date).bold = True


    def save_file(self):
        filename_tech = path + '//' + self.heading + ' -' + short_name + '.docx'
        self.doc.save(filename_tech)
        
        
    def create_doc(self):
        self.write_title()
        self.write_details()
        self.write_inside_addr()
        self.write_body()
        self.write_signature()
        self.save_file()


class tech(letter):
    heading = 'Letter of Technical Bid'
    def write_body(self):
        numbering_style = 'abcd'
        self.doc.add_paragraph('We, the undersigned, declare that:')
        self.doc.add_paragraph('We have examined and have no reservations to the Bidding Documents, including Addenda issued in accordance with Instructions to Bidders (ITB) Clause8.', style = numbering_style).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        name = self.doc.add_paragraph('We offer to execute in conformity with the Bidding Documents the following Works: ', style = numbering_style)
        name.add_run(Contract_ID).bold = True

        itb_18_1 = self.doc.add_paragraph('Our Bid consisting of the Technical Bid and the Price Bid shall be valid for a period of ', style = numbering_style)
        itb_18_1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        itb_18_1.add_run(days).bold = True
        itb_18_1.add_run(' from the date fixed for the bid submission deadline in accordance with the Bidding Documents, and it shall remain binding upon us and may be accepted at any time before the expiration of that period.')

        self.doc.add_paragraph('Our firm, including any subcontractors or suppliers for any part of the Contract, have nationalities from eligible countries in accordance with ITB 4.2 and meet the requirements of ITB 3.4, & 3.5', style = numbering_style).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        self.doc.add_paragraph('We are not participating, as a Bidder or as a subcontractor, in more than one Bid in this bidding process in accordance with ITB 4.3(e), other than alternative offers submitted in accordance with ITB 13.', style = numbering_style).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        self.doc.add_paragraph('Our firm, its affiliates or subsidiaries, including any Subcontractors or Suppliers for any part of the contract, has not been declared ineligible by DP, under the Employers country laws or official regulations or by an act of compliance with a decision of the United Nations Security Council;', style = numbering_style).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        self.doc.add_paragraph('We are not a government owned entity/We are a government owned entity but meet the requirements of ITB 4.5;', style = numbering_style).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        self.doc.add_paragraph('We declare that, we including any subcontractors or suppliers for any part of the contract do not have any conflict of interest in accordance with ITB 4.3 and we have not been punished for an offense relating to the concerned profession or business.', style = numbering_style).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        self.doc.add_paragraph('We declare that we are solely responsible for the authenticity of the documents submitted by us. The document and information submitted by us are true and correct. If any document/information given is found to be concealed at a later date, we shall accept any legal actions by the Employer.', style = numbering_style).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        self.doc.add_paragraph('We agree to permit the Employer/DP or its representative to inspect our accounts and records and other documents relating to the bid submission and to have them audited by auditors appointed by the Employer.', style = numbering_style).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        self.doc.add_paragraph('If our Bid is accepted, we commit to mobilizing key equipment and personnel in accordance with the requirements set forth in Section III (Evaluation and Qualification Criteria) and our technical proposal, or as otherwise agreed with the Employer.', style = numbering_style).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        credit = self.doc.add_paragraph('We are committed to submit the Letter of Commitment for Banks Undertaking for Line of Credit of ', style = numbering_style)
        credit.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        credit.add_run(line_of_credit).bold = True
        credit.add_run(' at the time of contract agreement, if the bid is awarded to us.')


class price(letter):
    heading = 'Letter of Price Bid'
    def write_body(self):
        numbering_style = 'abcd'
        self.doc.add_paragraph('We, the undersigned, declare that:')
        
        # List starts
        self.doc.add_paragraph('We have examined and have no reservations to the Bidding Documents, including Addenda issued in accordance with Instructions to Bidders (ITB) Clause8.', style = numbering_style).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        name = self.doc.add_paragraph('We offer to execute in conformity with the Bidding Documents the following Works: ', style = numbering_style)
        name.add_run(Contract_ID).bold = True

        self.doc.add_paragraph('The total price of our Bid, excluding any discounts offered in item (d) below is: NRs.……………………………; or when left blank is the Bid Price indicated in the Bill of Quantities.', style = numbering_style).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        self.doc.add_paragraph('The discounts offered and the methodology for their application are:………………………………', style = numbering_style).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        itb_18_1  = self.doc.add_paragraph('Our bid shall be valid for a period of ', style = numbering_style)
        itb_18_1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        itb_18_1.add_run(days).bold = True
        itb_18_1.add_run(' from the date fixed for the bid submission deadline in accordance with the Bidding Documents, and it shall remain binding upon us and may be accepted at any time before the expiration of that period;')

        self.doc.add_paragraph('If our bid is accepted, we commit to obtain a performance security in accordance with the Bidding Document;', style = numbering_style).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        self.doc.add_paragraph('We have paid, or will pay the following commissions, gratuities, or fees with respect to the bidding process or execution of the Contract:', style = numbering_style).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        none_none = self.doc.add_paragraph('             ')
        none_none.alignment = WD_ALIGN_PARAGRAPH.CENTER
        none_none.add_run('Name of Recipient').underline = True
        none_none.add_run('                     ')
        none_none.add_run('Address').underline = True
        none_none.add_run('	                ')
        none_none.add_run('Reason').underline = True
        none_none.add_run('	                ')
        none_none.add_run('Amount').underline = True

        none_none.add_run('\n                  ')
        none_none.add_run('None').italic = True
        none_none.add_run('			          ')
        none_none.add_run('None').italic = True
        none_none.add_run('		  ')
        none_none.add_run('None').italic = True
        none_none.add_run('		 ')
        none_none.add_run('None').italic = True

        self.doc.add_paragraph('We understand that this bid, together with your written acceptance thereof included in your notification of award, shall constitute a binding contract between us, until a formal contract is prepared and executed;', style = numbering_style).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        self.doc.add_paragraph('We understand that you are not bound to accept the lowest evaluated bid or any other bid that you may receive; and', style = numbering_style).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        self.doc.add_paragraph('We declare that we are solely responsible for the authenticity of the documents submitted by us.', style = numbering_style).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        self.doc.add_paragraph('We agree to permit the Employer/DP or its representative to inspect our accounts and records and other documents relating to the bid submission and to have them audited by auditors appointed by the Employer.', style = numbering_style).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY        


class bid(letter):
    heading = 'Letter of Bid'
    def write_body(self):
        numbering_style = 'abcd'
        self.doc.add_paragraph('We, the undersigned, declare that:')
        
        # List starts
        self.doc.add_paragraph('We have examined and have no reservations to the Bidding Documents, including Addenda issued in accordance with Instructions to Bidders (ITB) Clause8.', style = numbering_style).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        name = self.doc.add_paragraph('We offer to execute in conformity with the Bidding Documents the following Works: ', style = numbering_style)
        name.add_run(Contract_ID).bold = True

        self.doc.add_paragraph('The total price of our Bid, excluding any discounts offered in item (d) below is: NRs.……………………………; or when left blank is the Bid Price indicated in the Bill of Quantities.', style = numbering_style).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        self.doc.add_paragraph('The discounts offered and the methodology for their application are:………………………………', style = numbering_style).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        itb_18_1  = self.doc.add_paragraph('Our bid shall be valid for a period of ', style = numbering_style)
        itb_18_1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        itb_18_1.add_run(days).bold = True
        itb_18_1.add_run(' from the date fixed for the bid submission deadline in accordance with the Bidding Documents, and it shall remain binding upon us and may be accepted at any time before the expiration of that period;')

        self.doc.add_paragraph('If our bid is accepted, we commit to obtain a performance security in accordance with the Bidding Document;', style = numbering_style).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        self.doc.add_paragraph('Our firm, including any subcontractors or suppliers for any part of the Contract, have nationalities from eligible countries;', style = numbering_style).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        self.doc.add_paragraph('We, including any subcontractors or suppliers for any part of the contract, do not have any conflict of interest in accordance with ITB 4.3;', style = numbering_style).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        self.doc.add_paragraph('We are not participating, as a Bidder or as a subcontractor, in more than one Bid in this bidding process in accordance with ITB 4.3, other than alternative offers submitted in accordance with ITB 13.', style = numbering_style).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        self.doc.add_paragraph('Our firm, its affiliates or subsidiaries, including any Subcontractors or Suppliers for any part of the contract, has not been declared ineligible by DP, under the Employers country laws or official regulations or by an act of compliance with a decision of the United Nations Security Council;', style = numbering_style).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        self.doc.add_paragraph('We are not a government owned entity;', style = numbering_style).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        self.doc.add_paragraph('We understand that this bid, together with your written acceptance thereof included in your notification of award, shall constitute a binding contract between us, until a formal contract is prepared and executed;', style = numbering_style).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        self.doc.add_paragraph('We declare that, we have not been black listed as per ITB 3.4 and no conflict of interest in the proposed procurement proceedings and we have not been punished for an offense relating to the concerned profession or business.', style = numbering_style).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        self.doc.add_paragraph('We understand that you are not bound to accept the lowest evaluated bid or any other bid that you may receive; and', style = numbering_style).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        award = self.doc.add_paragraph('If awarded the contract, the person named below shall act as Contractor‘s Representative: ', style = numbering_style)
        award.add_run(person).bold = True

        self.doc.add_paragraph('We agree to permit the Employer/DP or its representative to inspect our accounts and records and other documents relating to the bid submission and to have them audited by auditors appointed by the Employer.', style = numbering_style).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY    


class declaration(letter):
    heading = 'Declaration by the Bidder'
    def write_body(self):
        temp = self.doc.add_paragraph('We hereby declare that we, ')
        temp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if jv:
            temp.add_run(jv + ' ' + firm_addr).bold = True
        else:
            temp.add_run(firm + ' ' + firm_addr).bold = True
        temp.add_run(' have gone through and understood the Bidding Document and we have prepared our Bid accordingly with signed and stamped in. We shall sign and stamp each page of contract agreement document in event of award of contract of us.')
        self.doc.add_paragraph('We further confirm that we have indicated price in schedule of Rates considering detailed description of item given in schedule of rates. We confirm that the rates quoted by us in schedule of rates include all activities in each item in relation with the performance the job.').alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        self.doc.add_paragraph('Moreover, we declare ourselves that our Firm has never been declared ineligible for the public procurement proceedings and do not have any conflicts of self interest in the proposed procurement proceeding and we have not been punished by an authority in the related profession or business till this date.').alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        temp = self.doc.add_paragraph('\n\n\n\n                                                                                                                                                       ')
        temp.add_run(person).bold = True
        if jv:
            temp.add_run('\n                                                                                                                                                      (Authorized JV Partner)').bold = True
        else:
            temp.add_run('\n                                                                                                                                                               (Proprietor)').bold = True
        self.doc.add_paragraph('\n                                                                                                                                                            Seal of the Firm')

    
    def create_doc(self):
        sections = self.doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        self.write_details()
        self.write_inside_addr()
        self.write_title(24)
        self.write_body()
        self.save_file()


if ls == '1':   #Small tender
    LoB = bid()
    LoB.create_doc()
    
if ls == '2':   #Large tender
    LoTB = tech()
    LoPB = price()
    LoTB.create_doc()
    LoPB.create_doc()
    
DbtB = declaration()
DbtB.create_doc()
