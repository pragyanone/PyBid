from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from .num2words import num2words
import os


class letter:
    def __init__(
        self,
        path,
        firm_name,
        firm_addr,
        jv_name,
        jv_addr,
        jv_repr,
        jv,
        short_name,
        contract_name,
        IoB,
        contract_ID,
        to,
        days,
        line_of_cr,
        person,
        date,
        jvList,
    ):
        self.path = path
        self.firm_name = firm_name
        self.firm_addr = firm_addr
        self.jv_name = jv_name
        self.jv_addr = jv_addr
        self.jv_repr = jv_repr
        self.jv = jv
        self.short_name = short_name
        self.contract_name = contract_name
        self.IoB = IoB
        self.contract_ID = contract_ID
        self.to = to
        self.days = days
        self.line_of_cr = line_of_cr
        self.person = person
        self.date = date
        self.jvList = jvList

        self.jvList = [line.split("; ") for line in self.jvList.split("\n")]
        self.jvList.sort()

        path_to_module = os.path.dirname(__file__)
        format_path = os.path.join(path_to_module, "PyTender-format.docx")
        self.doc = Document(format_path)
        self.designation = (
            "Managing Director" if "pvt" in self.firm_name.lower() else "Proprietor"
        )

    def write_title(self, size=16):
        title = self.doc.add_paragraph("\n\n")
        title.add_run(self.heading)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.runs[1].underline = title.runs[1].bold = True
        title.runs[1].font.size = Pt(size)

    def write_details(self):
        details = self.doc.add_paragraph("Date: ")
        details.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        details.add_run(self.date).bold = True
        details.add_run("\nName of the contract: ")
        details.add_run(self.contract_name).bold = True
        details.add_run("\nInvitation for Bid No.: ")
        details.add_run(self.IoB).bold = True
        details.add_run("\nContract Identification No.: ")
        details.add_run(self.contract_ID).bold = True

    def write_inside_addr(self):
        inside_addr = self.doc.add_paragraph("To:\n")
        inside_addr.alignment = WD_ALIGN_PARAGRAPH.LEFT
        inside_addr.add_run(self.to).bold = True

    def add_line(self, line):
        p = self.doc.add_paragraph(line, style="abcd")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        return p

    def write_signature(self):
        signature = self.doc.add_paragraph("\nName: ")
        if self.jv == 1:
            signature.add_run(self.person.rstrip()).bold = True
        elif self.jv == 2:
            signature.add_run(self.jv_repr.rstrip()).bold = True
        signature.add_run("\nIn the capacity of ")
        if self.jv == 2:
            signature.add_run("Authorized J.V. Representative").bold = True
        elif self.jv == 1:
            signature.add_run(self.designation).bold = True
        signature.add_run("\n\n\n\n\n\nSigned: ....................")
        signature.add_run("\nDuly authorized to sign the Bid for and on behalf of ")
        if self.jv_name:
            signature.add_run(self.jv_name).bold = True
        else:
            signature.add_run(self.firm_name).bold = True
        signature.add_run("\nDate: ")
        signature.add_run(self.date).bold = True

    def save_file(self):
        filename = (
            self.path + os.path.sep + self.heading + " -" + self.short_name + ".docx"
        )
        self.doc.save(filename)

    def create_doc(self):
        self.write_title()
        self.write_details()
        self.write_inside_addr()
        self.write_body()
        self.write_signature()
        self.save_file()


class tech(letter):
    heading = "Letter of Technical Bid"

    def write_body(self):
        self.doc.add_paragraph("We, the undersigned, declare that:")
        self.add_line(
            "We have examined and have no reservations to the Bidding Documents, including Addenda issued in accordance with Instructions to Bidders (ITB) Clause8."
        )

        name = self.add_line(
            "We offer to execute in conformity with the Bidding Documents the following Works: \n"
        )
        name.alignment = WD_ALIGN_PARAGRAPH.LEFT
        name.add_run(self.contract_ID).bold = True

        itb_18_1 = self.add_line(
            "Our Bid consisting of the Technical Bid and the Price Bid shall be valid for a period of "
        )
        itb_18_1.add_run(self.days).bold = True
        itb_18_1.add_run(
            " from the date fixed for the bid submission deadline in accordance with the Bidding Documents, and it shall remain binding upon us and may be accepted at any time before the expiration of that period."
        )

        self.add_line(
            "Our firm, including any subcontractors or suppliers for any part of the Contract, have nationalities from eligible countries in accordance with ITB 4.2 and meet the requirements of ITB 3.4, & 3.5"
        )
        self.add_line(
            "We are not participating, as a Bidder or as a subcontractor, in more than one Bid in this bidding process in accordance with ITB 4.3(e), other than alternative offers submitted in accordance with ITB 13."
        )
        self.add_line(
            "Our firm, its affiliates or subsidiaries, including any Subcontractors or Suppliers for any part of the contract, has not been declared ineligible by DP, under the Employers country laws or official regulations or by an act of compliance with a decision of the United Nations Security Council;"
        )
        self.add_line("We are not a government owned entity.")
        self.add_line(
            "We declare that, we including any subcontractors or suppliers for any part of the contract do not have any conflict of interest in accordance with ITB 4.3 and we have not been punished for an offense relating to the concerned profession or business."
        )
        self.add_line(
            "We declare that we are solely responsible for the authenticity of the documents submitted by us. The document and information submitted by us are true and correct. If any document/information given is found to be concealed at a later date, we shall accept any legal actions by the Employer."
        )
        self.add_line(
            "We agree to permit the Employer/DP or its representative to inspect our accounts and records and other documents relating to the bid submission and to have them audited by auditors appointed by the Employer."
        )
        self.add_line(
            "If our Bid is accepted, we commit to mobilizing key equipment and personnel in accordance with the requirements set forth in Section III (Evaluation and Qualification Criteria) and our technical proposal, or as otherwise agreed with the Employer."
        )

        credit = self.add_line(
            "We are committed to submit the Letter of Commitment for Banks Undertaking for Line of Credit of "
        )
        credit.add_run(self.line_of_cr).bold = True
        credit.add_run(
            " at the time of contract agreement, if the bid is awarded to us."
        )


class price(letter):
    heading = "Letter of Price Bid"

    def write_body(self):
        self.doc.add_paragraph("We, the undersigned, declare that:")

        # List starts
        self.add_line(
            "We have examined and have no reservations to the Bidding Documents, including Addenda issued in accordance with Instructions to Bidders (ITB) Clause8."
        )

        name = self.add_line(
            "We offer to execute in conformity with the Bidding Documents the following Works: \n"
        )
        name.alignment = WD_ALIGN_PARAGRAPH.LEFT
        name.add_run(self.contract_ID).bold = True

        self.add_line(
            "The total price of our Bid, excluding any discounts offered in item (d) below is: NRs.……………………………; or when left blank is the Bid Price indicated in the Bill of Quantities."
        )
        self.add_line(
            "The discounts offered and the methodology for their application are:………………………………"
        )

        itb_18_1 = self.add_line("Our bid shall be valid for a period of ")
        itb_18_1.add_run(self.days).bold = True
        itb_18_1.add_run(
            " from the date fixed for the bid submission deadline in accordance with the Bidding Documents, and it shall remain binding upon us and may be accepted at any time before the expiration of that period;"
        )

        self.add_line(
            "If our bid is accepted, we commit to obtain a performance security in accordance with the Bidding Document;"
        )
        self.add_line(
            "We have paid, or will pay the following commissions, gratuities, or fees with respect to the bidding process or execution of the Contract:"
        )

        none_none = self.doc.add_paragraph("             ")
        none_none.alignment = WD_ALIGN_PARAGRAPH.CENTER
        none_none.add_run("Name of Recipient").underline = True
        none_none.add_run("                           ")
        none_none.add_run("Address").underline = True
        none_none.add_run("	                  ")
        none_none.add_run("Reason").underline = True
        none_none.add_run("	                ")
        none_none.add_run("Amount").underline = True

        none_none.add_run("\n                  ")
        none_none.add_run("None").italic = True
        none_none.add_run("			          ")
        none_none.add_run("None").italic = True
        none_none.add_run("		  ")
        none_none.add_run("None").italic = True
        none_none.add_run("		 ")
        none_none.add_run("None").italic = True

        self.add_line(
            "We understand that this bid, together with your written acceptance thereof included in your notification of award, shall constitute a binding contract between us, until a formal contract is prepared and executed;"
        )
        self.add_line(
            "We understand that you are not bound to accept the lowest evaluated bid or any other bid that you may receive; and"
        )
        self.add_line(
            "We declare that we are solely responsible for the authenticity of the documents submitted by us."
        )
        self.add_line(
            "We agree to permit the Employer/DP or its representative to inspect our accounts and records and other documents relating to the bid submission and to have them audited by auditors appointed by the Employer."
        )


class bid(letter):
    heading = "Letter of Bid"

    def write_body(self):
        self.doc.add_paragraph("We, the undersigned, declare that:")

        # List starts
        self.add_line(
            "We have examined and have no reservations to the Bidding Documents, including Addenda issued in accordance with Instructions to Bidders (ITB) Clause8."
        )

        name = self.add_line(
            "We offer to execute in conformity with the Bidding Documents the following Works: "
        )
        name.alignment = WD_ALIGN_PARAGRAPH.LEFT
        name.add_run(self.contract_ID).bold = True

        self.add_line(
            "The total price of our Bid, excluding any discounts offered in item (d) below is: NRs.……………………………; or when left blank is the Bid Price indicated in the Bill of Quantities."
        )
        self.add_line(
            "The discounts offered and the methodology for their application are:………………………………"
        )

        itb_18_1 = self.add_line("Our bid shall be valid for a period of ")
        itb_18_1.add_run(self.days).bold = True
        itb_18_1.add_run(
            " from the date fixed for the bid submission deadline in accordance with the Bidding Documents, and it shall remain binding upon us and may be accepted at any time before the expiration of that period;"
        )

        self.add_line(
            "If our bid is accepted, we commit to obtain a performance security in accordance with the Bidding Document;"
        )
        self.add_line(
            "Our firm, including any subcontractors or suppliers for any part of the Contract, have nationalities from eligible countries;"
        )
        self.add_line(
            "We, including any subcontractors or suppliers for any part of the contract, do not have any conflict of interest in accordance with ITB 4.3;"
        )
        self.add_line(
            "We are not participating, as a Bidder or as a subcontractor, in more than one Bid in this bidding process in accordance with ITB 4.3, other than alternative offers submitted in accordance with ITB 13."
        )
        self.add_line(
            "Our firm, its affiliates or subsidiaries, including any Subcontractors or Suppliers for any part of the contract, has not been declared ineligible by DP, under the Employers country laws or official regulations or by an act of compliance with a decision of the United Nations Security Council;"
        )
        self.add_line("We are not a government owned entity;")
        self.add_line(
            "We understand that this bid, together with your written acceptance thereof included in your notification of award, shall constitute a binding contract between us, until a formal contract is prepared and executed;"
        )
        self.add_line(
            "We declare that, we have not been black listed as per ITB 3.4 and no conflict of interest in the proposed procurement proceedings and we have not been punished for an offense relating to the concerned profession or business."
        )
        self.add_line(
            "We understand that you are not bound to accept the lowest evaluated bid or any other bid that you may receive; and"
        )

        award = self.add_line(
            "If awarded the contract, the person named below shall act as Contractor‘s Representative: "
        )
        award.add_run(self.person).bold = True

        self.add_line(
            "We agree to permit the Employer/DP or its representative to inspect our accounts and records and other documents relating to the bid submission and to have them audited by auditors appointed by the Employer."
        )


class declaration(letter):
    heading = "Declaration by the Bidder"

    def write_body(self):
        temp = self.doc.add_paragraph("We hereby declare that we, ")
        temp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        temp.add_run(self.firm_name + ", " + self.firm_addr).bold = True
        temp.add_run(
            " have gone through and understood the Bidding Document and we have prepared our Bid accordingly with signed and stamped in. We shall sign and stamp each page of contract agreement document in event of award of contract of us."
        )
        self.doc.add_paragraph(
            "We further confirm that we have indicated price in schedule of Rates considering detailed description of item given in schedule of rates. We confirm that the rates quoted by us in schedule of rates include all activities in each item in relation with the performance the job."
        ).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        self.doc.add_paragraph(
            "Moreover, we declare ourselves that our Firm has never been declared ineligible for the public procurement proceedings and do not have any conflicts of self interest in the proposed procurement proceeding and we have not been punished by an authority in the related profession or business till this date."
        ).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        self.doc.add_paragraph("\n" * 4)
        temp = self.doc.add_paragraph()
        temp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        temp.add_run(self.person).bold = True
        temp.add_run(f"\n({self.designation})").bold = True
        temp.add_run("\nSeal of the Firm")

    def write_details(self):
        self.doc.add_paragraph("\n")
        letter.write_details(self)

    def create_doc(self, size):
        for section in self.doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        self.write_details()
        self.write_inside_addr()
        self.write_title(size)
        self.write_body()
        self.save_file()


class poa_self(declaration):
    heading = "Subject: Power of Attorney"

    def write_body(self):
        self.doc.add_paragraph("Dear Sir,")
        temp = self.doc.add_paragraph("With regards to the above subject I, ")
        temp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        temp.add_run(self.person).bold = True
        temp.add_run(', the Proprietor of this Firm "')
        temp.add_run(self.firm_name + ", " + self.firm_addr).bold = True
        temp.add_run(
            '" hereby inform you that as I, myself do all the activities of the firm, I appoint myself as the authorized representative to do all the official and site activities of the project mentioned above.'
        )

        temp = self.doc.add_paragraph("\n\n\n\n(Specimen Signature)\n")
        temp.add_run(self.person).bold = True

        temp = self.doc.add_paragraph()
        temp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        temp.add_run(self.person).bold = True
        temp.add_run("\n(" + self.designation + ")")

    def save_file(self):
        filename = (
            self.path + os.path.sep + "PoA Self" + " -" + self.short_name + ".docx"
        )
        self.doc.save(filename)


class jv_agr(letter):
    heading = "jvAgr"

    def write_body(self):
        temp = self.doc.paragraphs[0]
        temp.add_run("JOINT VENTURE AGREEMENT").underline = True
        temp.runs[0].bold = True
        temp = self.doc.add_paragraph("This AGREEMENT of JOINT VENTURE is made on ")
        temp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        temp.add_run(self.date).bold = True
        if len(self.jvList) > 2:
            temp.add_run(" among ")
        else:
            temp.add_run(" between ")

        for partner in self.jvList:
            temp.add_run(partner[3] + ", " + partner[4]).bold = True

            if partner[0] == "1":
                temp.add_run(' (hereinafter called the "Partner-In-Charge")')
            elif partner[0] == str(len(self.jvList)):
                temp.add_run(' (hereinafter called the "Second Partner(s)")')

            if partner[0] == str(len(self.jvList) - 1):
                temp.add_run(" and ")
            else:
                temp.add_run(", ")

        temp.add_run("of the Joint Venture.")

        temp = self.doc.add_paragraph("Whereas, ")
        temp.add_run(" ".join(self.to.split("\n")[1:])).bold = True
        temp.add_run(
            ' (hereinafter called the "The Employer") has invited the bid for the '
        )
        temp.add_run(self.contract_name).bold = True
        temp.add_run(", Contract No.: ")
        temp.add_run(self.contract_ID).bold = True
        temp.add_run(' (hereinafter called "The Works").\n')

        temp.add_run(
            'NOW WE UNDERSIGNED, responsible and authorized representatives of the "Joint Venture Partners" namely Partner-In-charge and Secondary Partner(s) DO AGREE as follows:'
        )

        temp = self.add_line("The name of joint venture will be ")
        temp.add_run(self.jv_name).bold = True
        temp.add_run(".")
        temp = self.add_line("The address of the Joint venture will be: ")
        temp.add_run(self.jv_addr).bold = True
        temp.add_run(".")

        temp = self.add_line("The Lead partner of the joint venture will be ")
        temp.add_run(self.jvList[0][3]).bold = True
        temp.add_run(" with its head office at ")
        temp.add_run(self.jvList[0][4]).bold = True
        temp.add_run(".")

        self.add_line(
            "The Lead Partner of joint venture is fully authorized to incur liabilities and receive instruction for on behalf of all partners and joint venture and the entire execution the contract including receiving of payment due. Both the partners of the joint venture shall be liable jointly severally for the execution of the contract. The lead partner shall also fully authorize to sign on behalf of the joint venture in all documents in regard the above mention project."
        )

        temp = self.add_line(
            "The specific responsibility of each party in the joint venture in brief would be as below "
        )
        temp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for partner in self.jvList:
            temp.add_run("\n")
            temp.add_run(partner[3]).bold = True
            temp.add_run(
                ": Execution/Administration/Technical expertise/Public relation/Plant & equipment,"
            )

        temp = self.add_line(
            "The required personnel shall be supplied by this joint venture to the project from internally agreed selection through the partner. Preference shall be given to the expertise available in the concerned field."
        )
        temp = self.add_line(
            "All the parties have agreed that the contribution of each party for this contract and the loss will also be shared as per each party contribution "
        )
        for partner in self.jvList:
            temp.add_run(
                partner[3] + "-" + partner[5] + "% [" + num2words(partner[5]) + "]"
            ).bold = True
            if partner[0] == str(len(self.jvList) - 1):
                temp.add_run(" and ")
            else:
                temp.add_run(", ")
        temp.add_run("of the total works.")

        self.add_line(
            "All the parties have agreed to make bid bond and required jointly/separately or any one party."
        )
        self.add_line(
            "All Parties will make the rates for preparation of tender document and the final rates to be quoted will be made and decided by board consisting of representatives form all parties."
        )
        self.add_line(
            "All  of the parties have agreed that a board consisting of a representative from each party will be formed which will be responsible for overall management to the job order to fulfill and the contractual obligation to complete the job within the stipulated time period."
        )
        self.add_line(
            "All of the parties have agreed to follow the instructions of the tender document and also the instruction of the consultant and the employer."
        )
        self.add_line(
            "All of the parties have agreed that the works will be up to the standard as per the specification as mentioned in the tender documents."
        )
        self.add_line(
            "All of the parties have agreed to open current account in the name of joint venture and will be operated by the representative from all parties."
        )
        temp = self.add_line(
            "All disputes relating to or arising out of the agreement will be sealed amicably between participating parties failing which shall be referred to arbitration as provided by Nepalese law and will be subjected to "
        )
        temp.add_run(self.jv_addr).bold = True
        temp.add_run(" for Justification.")

        self.add_line(
            "All the contract is awarded, this agreement will be in effect till the construction is completed in all respect and thereafter this agreement will null and void. If, otherwise the contact is not awarded, this agreement will be in effect till a decision by employer is made to award the contract, and thereafter will be null and void."
        )
        self.add_line(
            "This agreement written in English and executed in two (2) copies each of which shall be deemed as original and supersede all previous agreement understanding oral and written between all parties in respect to the matter hereof."
        )

        self.doc.add_paragraph(
            "The authorized representatives of the parties here-to have set their hands-on name and seal on the day-year first above written."
        )

    def create_table(self):
        table = self.doc.add_table(rows=2, cols=len(self.jvList))
        table.style = "CenteredTable"
        for i, partner in enumerate(self.jvList):
            if i > 0:
                table.rows[0].cells[i].text = "Secondary Partner"
            else:
                table.rows[0].cells[i].text = "Partner-In-Charge"

            txt = "\n" * 7
            for j in range(1, 5):
                txt += "\n" + partner[j]
            table.rows[1].cells[i].text = txt

    def create_doc(self):
        for section in self.doc.sections:
            section.top_margin = Cm(0.5)
            section.bottom_margin = Cm(0.3)
            section.left_margin = Cm(1)
            section.right_margin = Cm(1)
        self.write_body()
        self.create_table()
        self.save_file()


class poa(jv_agr):
    heading = "PoA"

    def write_date(self):
        temp = self.doc.paragraphs[0]
        temp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        temp.add_run("\n\n\n\n\n\nDate: ")
        temp.add_run(self.date).bold = True

    def write_body(self):
        temp = self.doc.add_paragraph("Sub: Power of Attorney")
        temp.bold = True
        temp.alignment = WD_ALIGN_PARAGRAPH.CENTER

        self.doc.add_paragraph("Dear Sir / Madam,")
        temp = self.doc.add_paragraph("We as the joint venture partners of ")
        temp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        temp.add_run(self.jv_name + ", " + self.jv_addr).bold = True
        temp.add_run(" hereby provide power of attorney to ")
        temp.add_run(self.jv_repr).bold = True
        temp.add_run(
            ', hereby called the "Authorized J.V. Representative" of this Joint Venture team to do and execute all of, and any of the acts and things as mentioned below in acceptance of all parties for the Execution of '
        )
        temp.add_run(self.contract_name).bold = True
        temp.add_run(", Contract ID No.: ")
        temp.add_run(self.contract_ID).bold = True
        temp.add_run(' (hereinafter called "The Works").')

        self.add_line(
            "To sign contract agreement document on behalf of the joint venture."
        )
        self.add_line(
            "To do all construction activities related with the contract, submit bills, and receive payment on behalf of the joint venture."
        )
        self.add_line(
            "To open a bank account and operate on behalf of the joint venture."
        )

        self.doc.add_paragraph(
            "As the joint venture hereby agree that all aforesaid acts, deeds and things done by the authorized jv representative shall be constructed as acts, deeds and things done by the joint venture and undertakes to rectify and conform all and whatsoever the authorized jv representative shall lawfully do the case to be done for joint venture of power hereby given."
        ).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        temp = self.doc.add_paragraph(
            "In witness whereof the Joint Venture has executed these deeds on "
        )
        temp.add_run(self.date).bold = True
        temp.add_run(".")

        temp = self.doc.add_paragraph("\n" * 4)
        temp.add_run(self.jv_repr).bold = True
        temp.add_run("\nSpecimen signature of Authorized J.V. Representative")

    def create_doc(self):
        for section in self.doc.sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)
        self.write_date()
        self.write_inside_addr()
        self.write_body()
        self.create_table()
        self.save_file()
