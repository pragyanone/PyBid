import docx
from docx import *
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH


doc = Document('PyTender-format.docx')
t1 = doc.add_table(rows=2, cols=2)
t1.style = 'CenteredTable'
t1.alignment = WD_TABLE_ALIGNMENT.CENTER
col1 = t1.columns[0].cells
col1
col1[0].text = 'Partner'


col1[1].text = str1

name, designation, company, address =
col2 = t1.columns[1].cells
col2[0].text = 'Secondary Partner'
#col2[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
col2[1].text = f'\n\n\n\n\n\n\n……………………………………..\n{name}\n{designation}\n{company}\n{address}'

doc.save('table.docx')